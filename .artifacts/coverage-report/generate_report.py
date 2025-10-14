from docx import Document
from docx.shared import Pt
import os

repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
workflow_path = os.path.join(repo_root, '.github', 'workflows', 'dotnet.yml')
csproj_path = os.path.join(repo_root, 'tests', 'TodoApi.Tests', 'TodoApi.Tests.csproj')
coverage_path = os.path.join(repo_root, 'tests', 'TodoApi.Tests', 'TestResults', '10160e50-49c6-46d4-ac07-992a49d678c4', 'coverage.cobertura.xml')

# Read files
with open(workflow_path, 'r', encoding='utf-8') as f:
    workflow = f.read()
with open(csproj_path, 'r', encoding='utf-8') as f:
    csproj = f.read()
with open(coverage_path, 'r', encoding='utf-8') as f:
    coverage = f.read()

# Create document
doc = Document()
doc.add_heading('SonarCloud Coverage Diagnosis Report', level=1)

p = doc.add_paragraph()
p.add_run('Repository: ').bold = True
p.add_run(os.path.basename(repo_root))

p = doc.add_paragraph()
p.add_run('Generated: ').bold = True
p.add_run(__import__('datetime').datetime.utcnow().isoformat() + 'Z')

# Findings summary
doc.add_heading('Findings summary', level=2)
doc.add_paragraph('SonarCloud shows 0% coverage likely because the Sonar scanner is configured to look for an OpenCover report at **TestResults/coverage.opencover.xml**, but the test run produced a Cobertura report (coverage.cobertura.xml).')

doc.add_heading('Files inspected', level=2)
for name, path in [('GitHub workflow', workflow_path), ('Test project', csproj_path), ('Coverage report', coverage_path)]:
    doc.add_paragraph(f'{name}: {path}')

# Add workflow
doc.add_heading('dotnet.yml (GitHub Actions Workflow)', level=2)
for line in workflow.splitlines():
    doc.add_paragraph(line, style='List Number')

# Add csproj
doc.add_heading('TodoApi.Tests.csproj', level=2)
for line in csproj.splitlines():
    doc.add_paragraph(line, style='List Number')

# Add coverage XML snippet
doc.add_heading('coverage.cobertura.xml (summary)', level=2)
# include only top section
for line in coverage.splitlines()[:120]:
    doc.add_paragraph(line, style='List Bullet')

# Detailed line-by-line notes
doc.add_heading('Detailed analysis and recommended fixes', level=2)
doc.add_paragraph('1) Issue: Workflow expects opencover file at **TestResults/coverage.opencover.xml** (sonar.cs.opencover.reportsPaths) but produced file is coverage.cobertura.xml.')
doc.add_paragraph('2) Fix (quick): Change workflow to pass Cobertura report to Sonar or produce an OpenCover report. SonarCloud accepts Cobertura for generic coverage but the property differs. For .NET, prefer OpenCover format by setting /p:CoverletOutputFormat=opencover and CoverletOutput to the path configured in sonar property.')

doc.add_paragraph('3) Example changes:')
doc.add_paragraph('- Keep sonar.cs.opencover.reportsPaths="**/TestResults/coverage.opencover.xml"')
doc.add_paragraph('- Ensure dotnet test uses: /p:CollectCoverage=true /p:CoverletOutputFormat=opencover /p:CoverletOutput=TestResults/coverage.opencover.xml')

doc.add_paragraph('4) Alternate: If you want to use Cobertura, update Sonar property to: "sonar.coverageReportPaths=**/TestResults/coverage.cobertura.xml" and pass the Cobertura file.')

doc.add_paragraph('5) Verify locally by running the dotnet test command and ensuring the expected coverage file is created under tests/TodoApi.Tests/TestResults/coverage.opencover.xml or coverage.cobertura.xml')

# Save
out_path = os.path.join(os.path.dirname(__file__), 'SonarCloud_Coverage_Report.docx')
doc.save(out_path)
print('Saved', out_path)
